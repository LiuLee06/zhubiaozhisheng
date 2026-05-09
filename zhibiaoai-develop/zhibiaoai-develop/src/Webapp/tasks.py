"""
标书处理工具函数 - 简化版
"""
from datetime import datetime
import os
import textwrap
from PyPDF2 import PdfReader
from docx import Document


def parse_pdf(file_path):
    """解析PDF文件"""
    with open(file_path, 'rb') as f:
        reader = PdfReader(f)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()


def parse_docx(file_path):
    """解析Word文档"""
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    return text.strip()


def generate_summary(full_text, max_length=300):
    """生成摘要"""
    if not full_text:
        return "文档内容为空，无法生成摘要"
    return textwrap.shorten(full_text, width=max_length, placeholder="...")


def export_to_word(bid_data, output_path):
    """将标书导出为Word文档"""
    doc = Document()
    doc.add_heading('投标文件', level=0)
    
    for section, content in bid_data.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    
    doc.save(output_path)
    return output_path


def process_uploaded_file(file_record):
    """处理上传的文件"""
    from Webapp.models import db
    
    file_record.status = '解析中'
    db.session.commit()
    
    if file_record.file_path.endswith('.pdf'):
        content = parse_pdf(file_record.file_path)
    elif file_record.file_path.endswith('.docx'):
        content = parse_docx(file_record.file_path)
    else:
        file_record.status = '不支持的文件格式'
        db.session.commit()
        raise ValueError('不支持的文件格式')
    
    summary = generate_summary(content)
    file_record.status = '解析完成'
    db.session.commit()
    
    return {
        'success': True,
        'message': f'文件 {file_record.file_name} 处理完成',
        'file_id': file_record.id,
        'content': content,
        'summary': summary
    }


def generate_bid(user_id, bid_data):
    """生成标书"""
    import requests
    from Webapp.config import Config
    
    print("进入 generate_bid 函数")
    
    # 配置豆包 API
    AI_API_KEY = Config.AI_API_KEY
    AI_BASE_URL = Config.AI_BASE_URL
    AI_MODEL = Config.AI_MODEL
    
    print(f"API配置 - Key: {AI_API_KEY[:10]}... URL: {AI_BASE_URL} Model: {AI_MODEL}")
    
    if not AI_API_KEY:
        raise ValueError("AI_API_KEY 未配置")
    
    # 生成标书内容
    result = {}
    sections = {
        "技术方案": bid_data.get('technical_requirements', ''),
        "实施计划": bid_data.get('implementation_plan', ''),
        "售后服务": bid_data.get('after_sales_service', '')
    }
    
    for section_name, section_content in sections.items():
        print(f"\n开始生成章节：{section_name}")
        
        prompt = f"""
根据以下信息生成{section_name}部分：

招标要求：
{bid_data.get('requirements', '')}

公司信息：
{bid_data.get('company_info', '')}

{section_name}要求：
{section_content}

请生成详细、专业的{section_name}内容，约500-800字。
"""
        
        try:
            # 调用豆包 API（添加超时控制）
            print(f"正在调用豆包 API...")
            response = requests.post(
                f"{AI_BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {AI_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": AI_MODEL,
                    "messages": [
                        {"role": "system", "content": "你是一个专业的标书撰写助手。"},
                        {"role": "user", "content": prompt}
                    ]
                },
                timeout=30  # 30秒超时
            )
            
            print(f"API 响应状态码：{response.status_code}")
            
            if response.status_code == 200:
                response_data = response.json()
                result[section_name] = response_data['choices'][0]['message']['content']
                print(f"章节 {section_name} 生成成功，长度：{len(result[section_name])}")
            else:
                error_msg = f"API返回错误 {response.status_code}: {response.text[:200]}"
                print(error_msg)
                result[section_name] = f"【{section_name}】\n\n本章节因API调用失败暂时无法生成。\n\n错误信息：{error_msg}"
                
        except requests.exceptions.Timeout:
            error_msg = "API调用超时"
            print(error_msg)
            result[section_name] = f"【{section_name}】\n\n本章节因API超时暂时无法生成。"
            
        except Exception as e:
            error_msg = f"调用API时发生错误：{str(e)}"
            print(error_msg)
            result[section_name] = f"【{section_name}】\n\n本章节生成时发生错误。\n\n错误信息：{error_msg}"
    
    print("\n所有章节生成完成，开始导出Word文档...")
    
    # 创建用户目录
    user_dir = os.path.join('generated_bids', str(user_id))
    os.makedirs(user_dir, exist_ok=True)
    print(f"用户目录：{user_dir}")
    
    # 生成文件名
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f"bid_{timestamp}.docx"
    output_path = os.path.join(user_dir, output_filename)
    print(f"输出路径：{output_path}")
    
    # 导出为Word文档
    try:
        export_to_word(result, output_path)
        print("Word文档导出成功！")
    except Exception as e:
        print(f"导出Word文档失败：{str(e)}")
        raise
    
    return {
        'success': True,
        'result': result,
        'document_path': output_path
    }